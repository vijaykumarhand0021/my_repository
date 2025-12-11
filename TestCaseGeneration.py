#main test genartion.py backup file its main file(Jyothi file)
from docx import Document
import csv
#from langchain.chat_models import ChatOpenAI
from langchain_openai import ChatOpenAI
import cv2
import re
import os
import cv2
#from moviepy import VideoFileClip
from moviepy.editor import VideoFileClip
import time
import base64
import pandas as pd

from langchain_community.callbacks import get_openai_callback
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
#from FRS import material_management, sales, production_planning

#from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field

class TestID_Name(BaseModel):
    """Date extraction from user input"""
    TestID: str = Field(default="", description="TestID Example: MV-01")
    TestName: str = Field(default="", description="One line short description. Example: \"Verification of Purchase Order Creation and Goods Receipt Process\"")

#from langchain_openai import ChatOpenAI
api_key = "sk-h2zLW0DZLXXgWkqwfvrZWg"
model = ChatOpenAI(
    base_url="https://genai-sharedservice-apac.pwcinternal.com",
    api_key=api_key,  
    model="azure.gpt-4o-2024-11-20",      
)
 
 # use your actual key


def process_video(video_path, seconds_per_frame=2):
    print("Entering process video")
    base64Frames = []
    base_video_path, _ = os.path.splitext(video_path)

    # Create a directory to save the frames
    frames_dir = f"{base_video_path}_frames1"
    os.makedirs(frames_dir, exist_ok=True)

    video = cv2.VideoCapture(video_path)
    total_frames = int(video.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = video.get(cv2.CAP_PROP_FPS)
    frames_to_skip = int(fps * seconds_per_frame)
    curr_frame=0

    # Loop through the video and extract frames at specified sampling rate
    frame_count = 0
    while curr_frame < total_frames - 1:
        video.set(cv2.CAP_PROP_POS_FRAMES, curr_frame)
        success, frame = video.read()
        if not success:
            break

        # Save frame as image
        frame_filename = os.path.join(frames_dir, f"frame_{frame_count:04d}.jpg")
        cv2.imwrite(frame_filename, frame)

        _, buffer = cv2.imencode(".jpg", frame)
        base64Frames.append(base64.b64encode(buffer).decode("utf-8"))
        curr_frame += frames_to_skip

        frame_count += 1

    video.release()

    print(f"Extracted {len(base64Frames)} frames")
    return base64Frames

def UpdateWordDocument_withST(csv_file, word_doc, updated_word_doc, sh, testcase_data):
    
    # df = pd.read_csv(csv_file, delimiter="|", engine="python")
    try:
        df = pd.read_csv(csv_file, delimiter="|", engine="python", encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(csv_file, delimiter="|", engine="python", encoding="latin1")

    data = df.values.tolist()

    # Load the Word document
    doc = Document(word_doc)
    

    table1 = doc.tables[0]
    cell_1 = table1.rows[0].cells[1]
    cell_1.text = testcase_data[0]

    cell_2 = table1.rows[1].cells[1]
    cell_2.text = testcase_data[1]

    # Access the first table
    table = doc.tables[1]

    tbl = table._tbl  # low-level table object
    last_row = table.rows[-1]._tr  # low-level row object
    tbl.remove(last_row)

    # Data to insert (only columns 0 to 2)
    new_data = data

    row = table.rows[1]
    row.cells[0].text = str(sh)

    cell = row.cells[0]
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(sh))
    run.bold = True
    

    # Update starting from second row (index 1)
    for i, row_data in enumerate(new_data, start=2):  # Start writing from row index 1
        if i >= len(table.rows):  # Add new rows if needed
            table.add_row()
        row = table.rows[i]
        for j, cell_value in enumerate(row_data):  # Only update first 3 columns
            row.cells[j].text = str(cell_value)

    # Save the updated document
    doc.save(updated_word_doc)
    print("Word Doc Table updated successfully!")
    return "Table updated successfully!"


def UpdateWordDocument(csv_file, word_doc, updated_word_doc):
    
    df = pd.read_csv(csv_file, delimiter="|", engine="python", quoting=csv.QUOTE_NONE)
    data = df.values.tolist()

    # Load the Word document
    doc = Document(word_doc)

    # Access the first table
    table = doc.tables[1]

    tbl = table._tbl  # low-level table object
    last_row = table.rows[-1]._tr  # low-level row object
    tbl.remove(last_row)

    # Data to insert (only columns 0 to 2)
    new_data = data

    # Update starting from second row (index 1)
    for i, row_data in enumerate(new_data, start=2):  # Start writing from row index 1
        if i >= len(table.rows):  # Add new rows if needed
            table.add_row()
        row = table.rows[i]
        for j, cell_value in enumerate(row_data):  # Only update first 3 columns
            row.cells[j].text = str(cell_value)

    # Save the updated document
    doc.save(updated_word_doc)
    print("Table updated successfully!")
    return "Table updated successfully!"


def split_list_with_overlap(lst, max_size=50, overlap=5):
    n = len(lst)
    if n <= max_size:
        return [lst]

    # Calculate number of chunks needed (ignore overlap for initial division)
    num_chunks = (n + (max_size - overlap - 1)) // (max_size - overlap)

    # Calculate base chunk size without overlap
    base_size = n // num_chunks
    remainder = n % num_chunks

    chunks = []
    start = 0

    for i in range(num_chunks):
        # Adjust chunk size to evenly distribute remainders
        chunk_size = base_size + (1 if i < remainder else 0)

        end = start + chunk_size

        # Add overlap to the start of next chunk (except the first)
        if i > 0:
            start -= overlap

        chunks.append(lst[start:end])
        start = end  # next chunk starts at end of current chunk

    return chunks

def ExtractCSV(data):
    csv_result = data.replace("“", "\"")
    csv_result = csv_result.replace("”", "\"")
    csv_pattern = "```csv\n(.*?)```"
    if "```csv" in data:
        csv_result = re.findall(csv_pattern, data, re.DOTALL)
        csv_data = csv_result[0]
    else:
        csv_data = data

    return csv_data

def CreateCSVFile(data, csv_file):

    csv_data = ExtractCSV(data)

    with open(csv_file, "w") as f:
        f.write(csv_data)

    return "CSV created successfully!"

def get_testid_name(test_cases):
    prompt = """Create one Test cases ID and Test Cases Name for the given Test Cases.

    Test Cases:
    {test_cases}

    Note: Generate only ID and Name in json
    """

    model1 = model.with_structured_output(TestID_Name)
    ai_message = model1.invoke(prompt)

    return [ai_message.TestID, ai_message.TestName]


def generateTestCasesVideo(base64Frames2, functional_spec):
    
    cost_details = {}
    cost_details["prompt_tokens"] = 0
    cost_details["completion_tokens"] = 0
    cost_details["total_tokens"] = 0
    cost_details["total_cost"] = 0
    test_cases = ""
    
    lst_of_frames = split_list_with_overlap(base64Frames2)
    print("Generating test cases for Video...")

    #For Video
    #Do not include variables/fields entered by user in test cases.
    #- Ensure logical grouping of video frames and generate complete test case scenario.
    #- Analyze which page or section comes after each page or section, and generate test cases and expected output correctly.
    #- For each test case scenario, positive test case should come after all negative test cases.
    #        - DO NOT generate incomplete test cases. 
    #Use words like Navigate, Dropdown, Select Tab etc. as appropriate.
    # Video frames are chunked and you'll be given each chunk of video frames in order to generate test cases. 
    # If it is not the first set of video frames, you'll be given test cases generated previously. 
    # Analyse and understand Video frames and functional specifications and generate test cases.
    # Generate final set of test cases using previously generated test cases along with new test cases from video frames and functional specifications.

    #You will also be given test cases generated from previous set of frames. You need to generate a final set of test cases using previous test cases and test cases generated from current set of frames. Do not miss any test cases.
    #- Include key terms, UI labels, on-screen messages, and error messages exactly as shown in the video frames.
    number_of_frames = len(lst_of_frames)
    for i in range(1, number_of_frames+1):
        messages=[
        {"role": "system", "content": f"""You are a QA test case generator. Your task is to analyze a software testing video composed of UI interactions and test scenarios. 
        The video has been broken down into frames captured every second. Each frame visually represents steps or parts of a test case. 
        There are total {len(base64Frames2)} frames. These frames are divided into {number_of_frames} sets with 5 frames overlapping in each set of frames. You will be given each set of frames to generate test cases. Do not speculate about missing parts of full video.
        Current frames - {i} out of {number_of_frames}

        You will be given functional specifications document which has the expected behaviour of the application. Use FSD as appropriately to generate test cases. Use T-codes and mandatory fields to check in test cases from FSD.
        Note: While generating test cases, mention the the field names only if the field name is mentioned in functional specifications document.
        
        Instructions:
        - 
        - Analyze and Understand key terms, UI labels, and what actions are being taken in each video frame.
        - Analyze the sequence of frames in chronological order, understand and reconstruct complete detailed test cases.  
        - Mention error and success messages that are displayed at bottom left of the screen if any.
        - Only include fields that are present in the Functional Specifications Document.
          Example: If "Material Code" is in the frame but not in FSD, do not include it. If "Purchase Order Number" is both in the frame and in FSD, include it.
        - For each test case, clearly mention how the user navigates to the page or section before performing the test action.  
          Example: Navigate to "Org. Data", enter valid Purchase Organisation, Purchase Group, and Company Code.
        - Observe the navigation and correctly mention navigation steps to the section or page. 
        - Mention correctly which page or section comes when you perform particular action.
        - Recognize and understand the short forms and abbreviations.
          Example: "Purch. Org." means "Purchase Organisation"
        - Generate meaningful and complete test cases. Test Description and Expected result should be meaningful.
          Examples for Expected Results:
          Example 1: User shall be able to view the "change process order: header-general data" screen
          Example 2: System shall be able to display the Process order creation screen
        - DO NOT miss any test cases. 
        - DO NOT mention anything which is not present in the video frames.
        - Do NOT generate any Field data. 
          Examples: 
          Incorrect Format:
          1. Test Description Input Information / Action: On the "Change Process Order: Header - General Data" screen, navigate to the "Goods Receipt" tab and update the "Stock Type" to "X Quality inspection".	
          2. Test Description Input Information / Action: Navigate to the "Dates/Quantities" tab to verify planned dates such as "Start: 31.05.2025" and "End: 31.05.2025".
          Correct Format:
          1. Test Description Input Information / Action: On the "Change Process Order: Header - General Data" screen, navigate to the "Goods Receipt" tab and update the "Stock Type".
          2. Test Description Input Information / Action: Navigate to the "Dates/Quantities" tab to verify planned dates.   
        - generate negative test cases.
        - DO NOT mention any keyboard shortcuts in test cases.
        
        
        Generate test cases in csv format, separated by "|".
        Example format:
        ```csv
        Example format:
        Step #|Test Description Input Information / Action|Expected Results
        1|On the "Create Process Order and Stock Transport Request" screen, enter the mandatory fields "Year" and "Month".|User shall be able to enter year and month.
        2|Click on the "Execute" button after entering "Year" and "Month".|System shall display the list of planned orders for the given input.
        ```
        Step #: Step number
        Test Description Input Information / Action: Include test steps in natural language
        Expected Results: Expected behaviour for the test case.
        """},
        {"role": "user", "content": [
            {"type": "text", "text": f"Functional Requirement Specifications (FSD): {functional_spec}"},
            {"type": "text", "text": "These are the frames from the video."},
            *map(lambda x: {"type": "image_url",
                            "image_url": {"url": f'data:image/jpg;base64,{x}', "detail": "high"}}, lst_of_frames[i-1])
            ],
        }
        ]
        
        with get_openai_callback() as cb:
            ai_message = model.invoke(messages)
            test_cases_csv = ai_message.content
            print(f"Set {i}:")
            print(test_cases_csv)
            test_cases += test_cases_csv
            print(cb.prompt_tokens)
            print(cb.completion_tokens)
            print(cb.total_tokens)
            cost_details["prompt_tokens"] += cb.prompt_tokens
            cost_details["completion_tokens"] += cb.completion_tokens
            cost_details["total_tokens"] += cb.total_tokens
            cost_details["total_cost"] += cb.total_cost
    
    test_cases_csv = test_cases
    
    #- Add token verification as a test case scenario. It is displayed after entering login details correctly. Add it as appropriately in order.     
    if len(lst_of_frames)>1:
        messages=[
        {"role": "system", "content": f"""You are a QA test case generator. 
        You will be given {number_of_frames} csv files generated from video frames in chronological order. You need to combine all csv data and modify as appropriate to generate final set of test cases.
        You will be given functional specifications as reference. Make use of it as appropriately while generating final set of test cases.
        
        Instructions:
        - If multiple test cases have same fields/ data and lead to same expected output, generate it as a single test case. Avoid duplication.
        - Ensure that step numbers are continuous.
        - Fix any wording inconsistencies or formatting issues.
        - Remove any obvious duplicates or overlapping cases, but do NOT remove valid negative/positive test pairs.
        - Do not miss any test cases. Do not generate incomplete test cases. 
        
        Generate test cases in csv format, separated by "|".
        Example format:
        ```csv
        Example format:
        Step #|Test Description Input Information / Action|Expected Results
        1|Log in to the VA system enter valid username and invalid password, click on the "I agree to the terms and conditions and privacy" check box, and click on the submit button.|User shall not be able to login into the system and system shall provide error message as "Login Failed Invalid username or password".
        ```

        Test Description Input Information / Action - Include test steps in natural language
        
        Note: Do not miss any test cases. All test cases are important."""},
        {"role": "user", "content": f"""Functional Requirement Specifications: {functional_spec}\n\n
         Here are the CSV files:\n {test_cases}""",
        }
        ]
        
        with get_openai_callback() as cb:
            ai_message = model.invoke(messages)
            test_cases_csv = ai_message.content
            print(cb.prompt_tokens)
            print(cb.completion_tokens)
            print(cb.total_tokens)
            cost_details["prompt_tokens"] += cb.prompt_tokens
            cost_details["completion_tokens"] += cb.completion_tokens
            cost_details["total_tokens"] += cb.total_tokens
            cost_details["total_cost"] += cb.total_cost
    
    
    
    print("Final set of test cases:\n", test_cases_csv)
    print("Test Cases generarted Successfully!")
    return test_cases_csv, cost_details

def start_process(path, video_path, word_doc_template, frs, sub_heading):
    import time
    start_time = time.time()
    folders = ["Dump", "Output"]
    for folder_path in folders:
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

    #base64Frames1 = process_video(video_path, seconds_per_frame=1.25)
    base64Frames1 = process_video(video_path, seconds_per_frame=0.5)
    response, cost_details = generateTestCasesVideo(base64Frames1, frs)
    testcases_fromVideo = ExtractCSV(response)
    csv_filenameVideo = "Dump/testcase_final_Video.csv"
    updated_word_doc_video = f"Output/{path}"
    resp = CreateCSVFile(testcases_fromVideo, csv_filenameVideo)

    testID_Name = get_testid_name(response)
    final_resp = UpdateWordDocument_withST(csv_filenameVideo, word_doc_template, updated_word_doc_video, sub_heading, testID_Name)
    
    # final_resp = UpdateWordDocument(csv_filenameVideo, word_doc_template, updated_word_doc_video)
    end_time = time.time()

    print("Time", end_time-start_time)
    print(cost_details)
    
    return "Test Cases generated!"

if __name__ == "__main__":
    #Material Management without field values

    #word_doc_template = r"C:\Users\jlolla001\Downloads\TestCases Generation\2. OQ Test script for Master  Final_Vendor_Blank 1.docx"
    word_doc_template = r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\test.docx"
    
    #frs_doc_path=r"C:\Users\vhande004\Downloads\TestCaseBackend 1\TestCaseBackend\Input_FRS\MaterialManagement.docx"
    
    
    #video_path =r"C:\Users\vhande004\Downloads\TestCaseBackend 1\TestCaseBackend\Input_Videos\Physical Inventory Cycle conunting process.mp4"
    #frs_doc_path=r"C:\Users\vhande004\Downloads\TestCaseBackend 1\TestCaseBackend\Input_FRS\EWM FRS for Physical inventory.docx"
    #video_path =r"C:\Users\vhande004\Downloads\TestCaseBackend 1\TestCaseBackend\Input_Videos\Direct process.mp4"
    #frs_doc_path=r"C:\Users\vhande004\Downloads\TestCaseBackend 1\TestCaseBackend\Input_FRS\Direct process FRS.docx"
    #===============================================================================================================================   
  
    frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\Goods Receipts process order.docx"
    video_path =r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\Recording (IM) 2025-12-04 194150.mp4"
  #-------------------------------------------------------------------------------------------------------------------------------  
    #video_path =r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\Resource, batch derivation sending and receiving.mp4"
    #frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\P2D resource, batch sending and receiving .docx"
   #---------------------------------------------------------------------------------------------------------------------------   
    #frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\EWM FRS for Physical inventory.docx"
    #video_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\Physical Inventory Cycle conunting process.mp4"
    #----------------------------------------------------------------------------------------------------------------------------------
    #frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\Direct process FRS.docx"
    #video_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\Direct process.mp4"  
    #-----------------------------------------------------------------------------------------------------
    ##frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\SDP(rem), DMR, Insppaln, Qualityinfo.docx"
    #video_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\SDP(rem), DMR, Insppaln, Qualityinfo.mp4"
    #---------------------------------------------------------------------------------------------------------
    
    #frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\MIC, INSPMETHD, SAMPRO, SAMPSCH, SAMPDRWPRC(SQRT).docx"
    #video_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\MIC, INSPMETHD, SAMPRO, SAMPSCH, SAMPDRWPRC(SQRT).mp4"
    #----------------------------------------------------------------------------------------------
    #frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\work center,equipment master,serilized equipment half.docx"
    #video_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\work center , equipment master , serilized equipment half Recording 2025-12-09 171544.mp4"
    #----------------------------------------------------------------------------------------------------------------------------------
    #frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\half serilized equipment, deletion funtionality,  functional location creation, general, technical tasklist half .docx"
    #video_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\half serilized equipment, deletion funttionality,  functional location creation, general, technical tasklist half Recording 2025-12-09 181307.mp4"
    #----------------------------------------------------------------------------------------------------------------------------------
    #frs_doc_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_FRS\O2C masters.docx"
    #video_path=r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\O2C Masters.mp4"
    
    
    
    output_filename = "O2C masters_Test_Cases_output.docx"
    sub_heading = 1
    #start_process(output_filename, video_path, word_doc_template, material_management)
    start_process(output_filename, video_path, word_doc_template, frs_doc_path,1)

   #video_path = r"C:\Users\jlolla001\Downloads\TestCases Generation\Videos\Videos\MaterialManagement.mp4"
   #video_path = r"C:\Users\vhande004\Downloads\ext\TestCaseBackend 1\TestCaseBackend\Input_Videos\MaterialManagement_testing.mp4"
   #video_path = r"C:\Users\jlolla001\Downloads\TestCases Generation\Videos\Sales&Distribution_QualityManagement.mp4"
   #video_path = r"C:\Users\jlolla001\Downloads\TestCases Generation\Videos\ProductionPlanning.mp4"
   #path = None